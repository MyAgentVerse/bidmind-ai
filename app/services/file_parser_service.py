"""File parsing service for extracting structured text from PDF and DOCX files.

Phase 1 of the BidMind AI deep-analysis upgrade.

This module preserves document structure that is critical for downstream
analysis and proposal generation:

  - Page numbers — so the AI can cite "per RFP § 4.2, p. 23"
  - Headings and section hierarchy — preserves the document outline
  - Tables — rendered as markdown so the LLM can read pricing schedules,
    compliance matrices, evaluation weights, etc.
  - Reading order — top-to-bottom, left-to-right
  - DOCX styles, headers/footers, and tables in document order
  - Optional OCR fallback for scanned PDFs

The legacy `parse_file()` API is preserved for backward compatibility:
it still returns ``Tuple[str, str]`` (text, file_type), but the text is now
a much richer markdown rendering with ``[Page N]`` markers, headings, and
tables. Existing call sites in ``uploads.py`` keep working unchanged.

A new API, ``parse_file_structured()``, returns a ``ParsedDocument`` containing
the rendered text AND a list of typed ``DocumentChunk`` objects for the
upcoming RAG / vector-store layer (Phase 2) and the compliance-matrix
builder.
"""

import logging
import statistics
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Optional, Tuple, Literal, Dict, Any

from app.utils.text_cleaning import normalize_text

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public data structures
# ---------------------------------------------------------------------------

ChunkType = Literal["heading", "paragraph", "table", "list", "header_footer"]


@dataclass
class DocumentChunk:
    """A single structural unit of a parsed document.

    Produced by the structured parser and consumed by:
      - chunk-aware analysis (Phase 1+)
      - the future vector store / RAG layer (Phase 2)
      - the future compliance-matrix builder (Phase 2)
    """

    text: str
    page: int                                      # 1-indexed (1 for DOCX)
    chunk_type: ChunkType
    section: Optional[str] = None                  # nearest preceding heading
    heading_level: Optional[int] = None            # 1-6 if chunk_type == "heading"
    bbox: Optional[Tuple[float, float, float, float]] = None  # PDF only

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class ParsedDocument:
    """Result of parsing a file."""

    full_text: str                 # Markdown-rendered, with [Page N] markers (PDF)
    chunks: List[DocumentChunk]
    file_type: str                 # "pdf" or "docx"
    page_count: int                # 1 for DOCX (no page concept)
    word_count: int
    char_count: int
    estimated_tokens: int          # Heuristic: char_count / 4
    has_tables: bool
    used_ocr: bool
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "full_text": self.full_text,
            "chunks": [c.to_dict() for c in self.chunks],
            "file_type": self.file_type,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "estimated_tokens": self.estimated_tokens,
            "has_tables": self.has_tables,
            "used_ocr": self.used_ocr,
            "warnings": list(self.warnings),
        }


# ---------------------------------------------------------------------------
# Tunables
# ---------------------------------------------------------------------------

# Heuristic: ~4 characters per token for English. Good enough for budget /
# chunking decisions. Swap in tiktoken later for exact counts.
CHARS_PER_TOKEN = 4

# If extracted text averages fewer than this many chars per page, assume
# the PDF is scanned and try OCR.
SCANNED_PDF_THRESHOLD_CHARS_PER_PAGE = 80

# A text block whose dominant font size is at least this multiple of the
# body font size is treated as a heading.
HEADING_FONT_RATIO = 1.15


# ---------------------------------------------------------------------------
# Service
# ---------------------------------------------------------------------------


class FileParserService:
    """Service for extracting structured text from uploaded documents."""

    # ---- Public API (backward compatible) -------------------------------

    @staticmethod
    def parse_file(file_path: str) -> Tuple[str, str]:
        """Parse a file and return ``(rich_markdown_text, file_type)``.

        Backward-compatible signature. The returned text now contains:
          - ``[Page N]`` markers (PDF only)
          - Markdown headings (``#``, ``##``, ...)
          - Markdown tables
          - Preserved reading order

        Existing call sites keep working unchanged; they just get a richer
        string back.
        """
        parsed = FileParserService.parse_file_structured(file_path)
        return parsed.full_text, parsed.file_type

    @staticmethod
    def parse_file_structured(file_path: str) -> ParsedDocument:
        """Parse a file and return a ``ParsedDocument`` with rich text + chunks."""
        path = Path(file_path)

        if not path.exists():
            raise IOError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        try:
            if extension == ".pdf":
                return FileParserService._parse_pdf(file_path)
            elif extension == ".docx":
                return FileParserService._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}")
            raise

    @staticmethod
    def validate_file_content(file_content: bytes, file_extension: str) -> bool:
        """Validate that file content matches the extension. (Unchanged API.)"""
        if not file_content or len(file_content) == 0:
            return False

        extension = file_extension.lower().lstrip(".")

        if extension == "pdf":
            return file_content.startswith(b"%PDF")
        elif extension == "docx":
            return file_content.startswith(b"PK")

        return True

    # ---- PDF parsing ----------------------------------------------------

    @staticmethod
    def _parse_pdf(file_path: str) -> ParsedDocument:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF is required for PDF parsing. "
                "Install with: pip install pymupdf"
            )

        logger.info(f"Parsing PDF (structured): {file_path}")

        chunks: List[DocumentChunk] = []
        rendered_pages: List[str] = []
        warnings: List[str] = []
        has_tables = False
        used_ocr = False

        doc = fitz.open(file_path)
        try:
            page_count = len(doc)

            # Pass 1: estimate the body font size by character-weighted mode.
            # Anything significantly larger is treated as a heading.
            body_font_size = FileParserService._estimate_body_font_size(doc)

            # State carried across pages so headings on page N apply to
            # paragraphs on page N+1 until the next heading.
            current_section: Optional[str] = None

            for page_index in range(page_count):
                page = doc[page_index]
                page_number = page_index + 1

                page_chunks, current_section = FileParserService._parse_pdf_page(
                    page,
                    page_number,
                    body_font_size,
                    current_section,
                )

                if any(c.chunk_type == "table" for c in page_chunks):
                    has_tables = True

                # Render this page as markdown for the legacy string output.
                page_md_lines = [f"[Page {page_number}]", ""]
                for c in page_chunks:
                    if c.chunk_type == "heading":
                        level = max(1, min(c.heading_level or 2, 6))
                        page_md_lines.append("#" * level + " " + c.text)
                    else:
                        # paragraphs and tables (already markdown)
                        page_md_lines.append(c.text)
                    page_md_lines.append("")

                rendered_pages.append("\n".join(page_md_lines))
                chunks.extend(page_chunks)

            full_text = "\n".join(rendered_pages)
            full_text = normalize_text(full_text)

            # OCR fallback for scanned PDFs
            avg_chars_per_page = (len(full_text) / page_count) if page_count else 0
            if avg_chars_per_page < SCANNED_PDF_THRESHOLD_CHARS_PER_PAGE:
                logger.warning(
                    f"PDF appears to be scanned "
                    f"({avg_chars_per_page:.0f} chars/page < "
                    f"{SCANNED_PDF_THRESHOLD_CHARS_PER_PAGE}). Attempting OCR."
                )
                ocr_text, ocr_chunks, ocr_succeeded = FileParserService._ocr_pdf(doc)
                if ocr_succeeded and len(ocr_text) > len(full_text):
                    full_text = normalize_text(ocr_text)
                    chunks = ocr_chunks
                    used_ocr = True
                else:
                    warnings.append(
                        "PDF appears to be scanned but OCR was unavailable "
                        "or produced no text. Install pytesseract + pillow "
                        "(and the tesseract binary) for scanned PDF support."
                    )
        finally:
            doc.close()

        word_count = len(full_text.split())
        char_count = len(full_text)
        estimated_tokens = char_count // CHARS_PER_TOKEN

        logger.info(
            f"PDF parsed: {page_count} pages, {char_count} chars, "
            f"~{estimated_tokens} tokens, {len(chunks)} chunks, "
            f"tables={has_tables}, ocr={used_ocr}"
        )

        return ParsedDocument(
            full_text=full_text,
            chunks=chunks,
            file_type="pdf",
            page_count=page_count,
            word_count=word_count,
            char_count=char_count,
            estimated_tokens=estimated_tokens,
            has_tables=has_tables,
            used_ocr=used_ocr,
            warnings=warnings,
        )

    @staticmethod
    def _estimate_body_font_size(doc) -> float:
        """Estimate the body font size by character-weighted mode.

        Walks up to 20 pages spread evenly across the document, sums the
        character count for each rounded font size, and returns the size
        with the most characters. That's the body font size.
        """
        size_to_chars: Dict[float, int] = {}
        page_count = len(doc)
        sample_indices = (
            list(range(page_count))
            if page_count <= 20
            else [int(i * page_count / 20) for i in range(20)]
        )

        for i in sample_indices:
            try:
                page_dict = doc[i].get_text("dict")
            except Exception:
                continue
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # 0 = text block
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        size = round(span.get("size", 10.0), 1)
                        text = span.get("text", "")
                        size_to_chars[size] = size_to_chars.get(size, 0) + len(text)

        if not size_to_chars:
            return 10.0

        body_size = max(size_to_chars.items(), key=lambda kv: kv[1])[0]
        logger.debug(f"Estimated body font size: {body_size}")
        return body_size

    @staticmethod
    def _parse_pdf_page(
        page,
        page_number: int,
        body_font_size: float,
        current_section: Optional[str],
    ) -> Tuple[List[DocumentChunk], Optional[str]]:
        """Parse a single PDF page into chunks.

        Returns ``(chunks, updated_section)``. The updated section is the
        most recent heading seen on this page (or the previous one if no
        heading was found).
        """
        # 1. Find tables and remember their bboxes so we can suppress
        #    duplicated text from the regular block extraction.
        table_items: List[Tuple[Tuple[float, float, float, float], str]] = []
        try:
            tables = page.find_tables()
            for tbl in tables:
                try:
                    rows = tbl.extract()
                    if not rows:
                        continue
                    md = FileParserService._render_table_as_markdown(rows)
                    if md:
                        table_items.append((tuple(tbl.bbox), md))
                except Exception as e:
                    logger.debug(f"Skipping a table on page {page_number}: {e}")
        except Exception as e:
            logger.debug(f"find_tables() failed on page {page_number}: {e}")

        # 2. Walk text blocks in reading order, skipping anything inside a table.
        try:
            page_dict = page.get_text("dict")
        except Exception as e:
            logger.warning(f"get_text(dict) failed on page {page_number}: {e}")
            page_dict = {"blocks": []}

        # (bbox, chunk) pairs so we can sort by reading order at the end
        items: List[Tuple[Tuple[float, float, float, float], DocumentChunk]] = []

        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:  # not a text block
                continue

            bbox = tuple(block.get("bbox", (0, 0, 0, 0)))

            # Skip blocks that fall inside a detected table
            if any(_bbox_inside(bbox, tbb) for tbb, _ in table_items):
                continue

            # Combine spans within the block, tracking dominant font size
            block_text_parts: List[str] = []
            block_sizes: List[float] = []
            for line in block.get("lines", []):
                line_parts = []
                for span in line.get("spans", []):
                    txt = span.get("text", "")
                    if txt:
                        line_parts.append(txt)
                        block_sizes.append(span.get("size", body_font_size))
                if line_parts:
                    block_text_parts.append(" ".join(line_parts).strip())

            block_text = "\n".join(block_text_parts).strip()
            if not block_text:
                continue

            dominant_size = (
                statistics.median(block_sizes) if block_sizes else body_font_size
            )

            # Heading detection: significantly larger font, short, single line
            is_heading = (
                dominant_size >= body_font_size * HEADING_FONT_RATIO
                and len(block_text) <= 200
                and "\n" not in block_text
            )

            if is_heading:
                ratio = dominant_size / body_font_size if body_font_size else 1.0
                if ratio >= 1.6:
                    level = 1
                elif ratio >= 1.4:
                    level = 2
                elif ratio >= 1.25:
                    level = 3
                else:
                    level = 4
                chunk = DocumentChunk(
                    text=block_text,
                    page=page_number,
                    chunk_type="heading",
                    heading_level=level,
                    section=block_text,
                    bbox=bbox,
                )
                current_section = block_text
            else:
                chunk = DocumentChunk(
                    text=block_text,
                    page=page_number,
                    chunk_type="paragraph",
                    section=current_section,
                    bbox=bbox,
                )
            items.append((bbox, chunk))

        # 3. Add tables as chunks at their own positions
        for tbb, md in table_items:
            items.append(
                (
                    tbb,
                    DocumentChunk(
                        text=md,
                        page=page_number,
                        chunk_type="table",
                        section=current_section,
                        bbox=tbb,
                    ),
                )
            )

        # 4. Reading-order sort: top-to-bottom, then left-to-right.
        #    Round y to absorb sub-line jitter so columns sort correctly.
        items.sort(key=lambda item: (round(item[0][1], 0), round(item[0][0], 0)))
        chunks = [c for _, c in items]

        return chunks, current_section

    @staticmethod
    def _render_table_as_markdown(rows: List[List[Optional[str]]]) -> str:
        """Render a list-of-rows table as a GitHub-flavored markdown table."""
        cleaned_rows: List[List[str]] = []
        for row in rows:
            cleaned = [
                (cell or "").replace("\n", " ").replace("|", "/").strip()
                for cell in row
            ]
            if any(cell for cell in cleaned):
                cleaned_rows.append(cleaned)

        if not cleaned_rows:
            return ""

        max_cols = max(len(r) for r in cleaned_rows)
        for r in cleaned_rows:
            while len(r) < max_cols:
                r.append("")

        header = cleaned_rows[0]
        body = cleaned_rows[1:] if len(cleaned_rows) > 1 else []

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * max_cols) + " |",
        ]
        for row in body:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    @staticmethod
    def _ocr_pdf(doc) -> Tuple[str, List[DocumentChunk], bool]:
        """Best-effort OCR using pytesseract.

        Returns ``(rendered_text, chunks, success)``. If pytesseract or
        Pillow are not installed, returns ``("", [], False)`` and the caller
        falls back to whatever it had.
        """
        try:
            import pytesseract
            from PIL import Image
            import io
        except ImportError:
            logger.info(
                "OCR not attempted: pytesseract / Pillow not installed. "
                "To enable scanned-PDF support: "
                "pip install pytesseract pillow "
                "(and install the tesseract binary on your system)."
            )
            return "", [], False

        rendered_pages: List[str] = []
        chunks: List[DocumentChunk] = []
        for page_index in range(len(doc)):
            page = doc[page_index]
            try:
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                page_text = pytesseract.image_to_string(img)
            except Exception as e:
                logger.warning(f"OCR failed on page {page_index + 1}: {e}")
                page_text = ""

            page_text = page_text.strip()
            if page_text:
                rendered_pages.append(f"[Page {page_index + 1}]\n\n{page_text}")
                chunks.append(
                    DocumentChunk(
                        text=page_text,
                        page=page_index + 1,
                        chunk_type="paragraph",
                    )
                )

        return "\n\n".join(rendered_pages), chunks, bool(rendered_pages)

    # ---- DOCX parsing ---------------------------------------------------

    @staticmethod
    def _parse_docx(file_path: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )

        logger.info(f"Parsing DOCX (structured): {file_path}")

        doc = Document(file_path)
        chunks: List[DocumentChunk] = []
        md_lines: List[str] = []
        has_tables = False
        current_section: Optional[str] = None

        # Walk the body in document order, mixing paragraphs and tables.
        # python-docx doesn't expose ordered iteration directly, so we walk
        # the underlying lxml children and look up wrappers by element id.
        body = doc.element.body
        para_by_id = {id(p._element): p for p in doc.paragraphs}
        table_by_id = {id(t._element): t for t in doc.tables}

        for child in body.iterchildren():
            child_id = id(child)

            if child_id in para_by_id:
                para = para_by_id[child_id]
                text = (para.text or "").strip()
                if not text:
                    continue

                style_name = (para.style.name if para.style else "") or ""
                heading_level = FileParserService._docx_heading_level(style_name)

                if heading_level:
                    chunks.append(
                        DocumentChunk(
                            text=text,
                            page=1,
                            chunk_type="heading",
                            heading_level=heading_level,
                            section=text,
                        )
                    )
                    current_section = text
                    md_lines.append("#" * heading_level + " " + text)
                else:
                    chunks.append(
                        DocumentChunk(
                            text=text,
                            page=1,
                            chunk_type="paragraph",
                            section=current_section,
                        )
                    )
                    md_lines.append(text)
                md_lines.append("")

            elif child_id in table_by_id:
                table = table_by_id[child_id]
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                md = FileParserService._render_table_as_markdown(rows)
                if md:
                    has_tables = True
                    chunks.append(
                        DocumentChunk(
                            text=md,
                            page=1,
                            chunk_type="table",
                            section=current_section,
                        )
                    )
                    md_lines.append(md)
                    md_lines.append("")

        # Headers and footers (collect first occurrence of each unique block)
        seen_hf: set = set()
        for section in doc.sections:
            for hf, label in (
                (section.header, "header"),
                (section.footer, "footer"),
            ):
                try:
                    txt = "\n".join(
                        p.text.strip() for p in hf.paragraphs if p.text.strip()
                    )
                except Exception:
                    txt = ""
                if txt and txt not in seen_hf:
                    seen_hf.add(txt)
                    chunks.append(
                        DocumentChunk(
                            text=f"[{label}] {txt}",
                            page=1,
                            chunk_type="header_footer",
                        )
                    )
                    md_lines.append(f"_{label}: {txt}_")
                    md_lines.append("")

        full_text = normalize_text("\n".join(md_lines))
        word_count = len(full_text.split())
        char_count = len(full_text)
        estimated_tokens = char_count // CHARS_PER_TOKEN

        logger.info(
            f"DOCX parsed: {char_count} chars, ~{estimated_tokens} tokens, "
            f"{len(chunks)} chunks, tables={has_tables}"
        )

        return ParsedDocument(
            full_text=full_text,
            chunks=chunks,
            file_type="docx",
            page_count=1,  # python-docx has no page concept
            word_count=word_count,
            char_count=char_count,
            estimated_tokens=estimated_tokens,
            has_tables=has_tables,
            used_ocr=False,
            warnings=[],
        )

    @staticmethod
    def _docx_heading_level(style_name: str) -> Optional[int]:
        """Map a python-docx paragraph style name to a heading level (1-6)."""
        if not style_name:
            return None
        s = style_name.strip().lower()
        if s == "title":
            return 1
        if s.startswith("heading "):
            try:
                level = int(s.split(" ", 1)[1])
                return max(1, min(level, 6))
            except (ValueError, IndexError):
                return None
        return None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _bbox_inside(
    inner: Tuple[float, float, float, float],
    outer: Tuple[float, float, float, float],
    margin: float = 2.0,
) -> bool:
    """Return True if ``inner`` bbox is (mostly) inside ``outer`` bbox."""
    if not inner or not outer:
        return False
    ix0, iy0, ix1, iy1 = inner
    ox0, oy0, ox1, oy1 = outer
    return (
        ix0 >= ox0 - margin
        and iy0 >= oy0 - margin
        and ix1 <= ox1 + margin
        and iy1 <= oy1 + margin
    )
